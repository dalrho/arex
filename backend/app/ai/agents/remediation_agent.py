import os
import uuid
import logging
import difflib
from datetime import datetime, timezone
from typing import Dict, Any, List
from pydantic import BaseModel

from app.core.dependencies import SessionLocal
from app.models.document import Document
from app.models.regulation_update import RegulationUpdate
from app.models.remediation_draft import RemediationDraft
from app.ai.llm_client import llm_client
from app.ai.tools.citation_tool import validate_citations

logger = logging.getLogger("arex.remediation-agent")

class RemediationAgentOutput(BaseModel):
    proposed_text: str
    citations: List[str]
    rationale: str

def generate_diff_metadata(original: str, proposed: str) -> Dict[str, List[str]]:
    """
    Utility using python difflib to compute structured line difference metadata.
    """
    orig_lines = original.splitlines()
    prop_lines = proposed.splitlines()
    
    diff = list(difflib.ndiff(orig_lines, prop_lines))
    
    added = []
    removed = []
    for line in diff:
        if line.startswith("+ "):
            added.append(line[2:])
        elif line.startswith("- "):
            removed.append(line[2:])
            
    return {
        "added": added,
        "removed": removed
    }


def _generate_proposed_new_sop(
    db,
    regulation: RegulationUpdate,
    organization_id_str: str,
) -> List[str]:
    """
    Create or update a Proposed New SOP draft when there are no usable matched documents.
    Returns a list of remediation draft ID strings.
    """
    import time
    from app.core.profiler import RequestProfiler
    import docx

    draft_ids: List[str] = []
    safe_title = "".join(c for c in regulation.title if c.isalnum() or c in (" ", "-", "_")).strip()
    filename = f"Proposed New SOP - {safe_title}.docx"

    existing_draft = db.query(RemediationDraft).join(Document, RemediationDraft.sop_id == Document.id).filter(
        RemediationDraft.regulation_id == regulation.id,
        Document.filename.like("Proposed New SOP - %")
    ).first()

    if existing_draft:
        logger.info(f"Proposed new SOP draft already exists for regulation {regulation.id}. Re-generating draft content.")
        new_doc = db.query(Document).filter(Document.id == existing_draft.sop_id).first()
    else:
        unique_file_id = uuid.uuid4()
        file_path = f"/app/storage/{unique_file_id}.docx"

        doc_obj = docx.Document()
        doc_obj.add_paragraph("Proposed New SOP Placeholder")
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        doc_obj.save(file_path)

        new_doc = Document(
            id=unique_file_id,
            organization_id=uuid.UUID(organization_id_str),
            filename=filename,
            file_path=file_path,
            version=1,
            parsed_text="[Proposed New SOP Placeholder]",
            created_at=datetime.now(timezone.utc)
        )
        t_db = time.time()
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)
        RequestProfiler.log_metric("database_write_time", time.time() - t_db)

    system_prompt = (
        "You are an expert Compliance Officer and GxP technical writer. "
        "Write a complete, professional, standalone Standard Operating Procedure (SOP) "
        "from scratch in response to the provided FDA regulation update. "
        "The document MUST be GxP-compliant, structured, and complete. "
        "Do not include snippets, comments, or summaries — output the full, complete document text."
    )

    user_prompt = (
        f"Write a complete, compliance SOP from scratch for: {filename}\n\n"
        f"FDA Regulation Update:\n"
        f"Title: {regulation.title}\n"
        f"Content Summary:\n{regulation.summary or regulation.raw_content}\n\n"
        f"The text should have standard sections: 1.0 Purpose, 2.0 Scope, 3.0 Responsibilities, 4.0 Procedure, etc."
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]

    logger.info("Invoking LLM to generate Proposed New SOP from scratch...")
    try:
        result = llm_client.get_completion(
            messages=messages,
            response_model=RemediationAgentOutput,
            temperature=0.0
        )
    except Exception as llm_err:
        logger.error(f"LLM call failed for new SOP generation: {llm_err}. Using stub fallback.")
        result = RemediationAgentOutput(
            proposed_text=(
                f"1.0 Purpose\n"
                f"This SOP establishes procedures to comply with: {regulation.title}.\n\n"
                f"2.0 Scope\n"
                f"Applies to all personnel involved in relevant operations.\n\n"
                f"3.0 Responsibilities\n"
                f"Quality Assurance is responsible for implementation and training.\n\n"
                f"4.0 Procedure\n"
                f"4.1 Review the requirements of {regulation.title}.\n"
                f"4.2 Update affected processes to align with regulation requirements.\n"
                f"4.3 Document all changes and obtain QA approval.\n\n"
                f"5.0 References\n"
                f"- {regulation.title}\n\n"
                f"[NOTE: This is an AI-generated stub. The LLM API was unavailable (quota exceeded). "
                f"Please edit this draft with the full SOP content.]"
            ),
            citations=[regulation.title],
            rationale=f"Stub draft generated due to LLM API unavailability. Manual review required."
        )

    if existing_draft:
        existing_draft.proposed_revision = result.proposed_text
        existing_draft.diff_content = {"added": result.proposed_text.splitlines(), "removed": []}
        existing_draft.explanation = result.rationale
        existing_draft.status = "Draft"
        t_db = time.time()
        db.commit()
        db.refresh(existing_draft)
        RequestProfiler.log_metric("database_write_time", time.time() - t_db)
        draft_ids.append(str(existing_draft.id))
    else:
        new_draft = RemediationDraft(
            id=uuid.uuid4(),
            sop_id=new_doc.id,
            regulation_id=regulation.id,
            proposed_revision=result.proposed_text,
            current_content="",
            diff_content={"added": result.proposed_text.splitlines(), "removed": []},
            explanation=result.rationale,
            status="Draft"
        )
        t_db = time.time()
        db.add(new_draft)
        db.commit()
        db.refresh(new_draft)
        RequestProfiler.log_metric("database_write_time", time.time() - t_db)
        draft_ids.append(str(new_draft.id))

    logger.info(f"Remediation Agent successfully generated proposed new SOP draft: {draft_ids[0]}")
    return draft_ids


def _process_single_document(
    doc_id_val,
    regulation_id_str: str,
    organization_id_str: str,
    system_prompt: str
):
    """
    Worker function to process a single document compliance remediation.
    Runs inside a thread pool with its own database session.
    """
    import time
    import uuid
    from app.core.dependencies import SessionLocal
    from app.core.profiler import RequestProfiler
    from app.models.document import Document
    from app.models.regulation_update import RegulationUpdate
    from app.models.remediation_draft import RemediationDraft
    from app.services.embeddings.embedding_service import embedding_service
    from app.services.vector_db.qdrant_client import vector_db_client

    # Initialize thread-local profiler metrics
    RequestProfiler.reset()
    
    db_thread = SessionLocal()
    try:
        doc_id = uuid.UUID(str(doc_id_val))
        doc = db_thread.query(Document).filter(Document.id == doc_id).first()
        if not doc or not doc.parsed_text:
            logger.warning(f"SOP Document {doc_id} has empty content, skipping in thread.")
            return None, RequestProfiler.get_metrics()

        # Fetch regulation in this thread session
        regulation = db_thread.query(RegulationUpdate).filter(
            RegulationUpdate.id == uuid.UUID(regulation_id_str)
        ).first()
        if not regulation:
            logger.error(f"Regulation {regulation_id_str} not found in database for thread.")
            return None, RequestProfiler.get_metrics()

        # RAG retrieval
        context_docs = []
        try:
            query_text = f"{regulation.raw_content} {doc.parsed_text[:500]}"
            query_vector = embedding_service.get_embedding(query_text)
            org_id = uuid.UUID(organization_id_str)
            chunks = vector_db_client.search_chunks(
                query_vector=query_vector,
                organization_id=org_id,
                limit=5,
            )
            context_docs = [
                {
                    "document_name": c.get("document_id", "KB Document"),
                    "text_snippet": c.get("text", ""),
                    "confidence_score": round(c.get("score", 0.0) * 100, 1),
                }
                for c in chunks
                if c.get("score", 0.0) >= 0.5
            ]
            logger.info(
                f"[RemediationAgent Thread] RAG retrieved {len(context_docs)} chunks "
                f"for document '{doc.filename}'."
            )
        except Exception as rag_err:
            logger.warning(
                f"[RemediationAgent Thread] RAG retrieval failed for '{doc.filename}': {rag_err}."
            )

        user_prompt = (
            f"SOURCE REGULATION UPDATE:\n{regulation.raw_content}\n\n"
            f"SOP TARGET DOCUMENT ({doc.filename}):\n{doc.parsed_text}\n\n"
            "Review the target document and suggest changes required for compliance."
        )

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

        logger.info(
            f"[RemediationAgent Thread] Calling LLM | document='{doc.filename}' | "
            f"context_docs={len(context_docs)}"
        )
        attempts = 3
        result = None
        for attempt in range(1, attempts + 1):
            try:
                result = llm_client.get_completion(
                    messages=messages,
                    response_model=RemediationAgentOutput,
                    temperature=0.0,
                    context_docs=context_docs if context_docs else None,
                )
                
                is_valid = validate_citations(
                    citations=result.citations,
                    regulation_content=regulation.raw_content,
                    parsed_sections=regulation.parsed_sections
                )
                
                if is_valid:
                    break
                else:
                    logger.warning(f"Citations validation failed on attempt {attempt} for '{doc.filename}'.")
                    messages.append({
                        "role": "assistant",
                        "content": f"Proposed revisions:\n{result.proposed_text}\nCitations:\n{result.citations}"
                    })
                    messages.append({
                        "role": "user",
                        "content": "Your citations list cannot be empty and must map directly to the clauses in the regulation. Please retry."
                    })
            except Exception as ex:
                logger.error(f"Error on LLM call attempt {attempt} for '{doc.filename}': {ex}")
                if attempt == attempts:
                    logger.warning(f"All {attempts} LLM attempts failed. Fallback stub generated.")
                    result = RemediationAgentOutput(
                        proposed_text=(
                            f"{doc.parsed_text}\n\n"
                            f"--- COMPLIANCE UPDATE REQUIRED ---\n"
                            f"This SOP requires revision to comply with: {regulation.title}.\n"
                            f"[NOTE: This is an AI-generated stub. The LLM API was unavailable (quota exceeded).]"
                        ),
                        citations=[regulation.title],
                        rationale=f"Stub draft generated for {regulation.title} due to LLM error."
                    )
                    break

        if not result:
            return None, RequestProfiler.get_metrics()

        diff_meta = generate_diff_metadata(doc.parsed_text, result.proposed_text)

        existing_draft = db_thread.query(RemediationDraft).filter(
            RemediationDraft.sop_id == doc_id,
            RemediationDraft.regulation_id == regulation.id
        ).first()

        draft_id_out = None
        t_db = time.time()
        if existing_draft:
            logger.info(f"Draft already exists for {doc.filename}. Updating proposed text.")
            existing_draft.proposed_revision = result.proposed_text
            existing_draft.diff_content = diff_meta
            existing_draft.explanation = result.rationale
            existing_draft.status = "Draft"
            db_thread.commit()
            db_thread.refresh(existing_draft)
            draft_id_out = str(existing_draft.id)
        else:
            new_draft = RemediationDraft(
                id=uuid.uuid4(),
                sop_id=doc_id,
                regulation_id=regulation.id,
                proposed_revision=result.proposed_text,
                current_content=doc.parsed_text,
                diff_content=diff_meta,
                explanation=result.rationale,
                status="Draft"
            )
            db_thread.add(new_draft)
            db_thread.commit()
            db_thread.refresh(new_draft)
            draft_id_out = str(new_draft.id)
        
        RequestProfiler.log_metric("database_write_time", time.time() - t_db)
        return draft_id_out, RequestProfiler.get_metrics()
    except Exception as err:
        db_thread.rollback()
        logger.error(f"Failed processing document {doc_id_val} in thread: {err}")
        return None, RequestProfiler.get_metrics()
    finally:
        db_thread.close()


def run_remediation_agent(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    LangGraph agent node that reviews impacted SOP documents and proposes draft revisions.
    Performs validation on citations and computes side-by-side diff details.
    """
    from app.core.profiler import RequestProfiler
    import time
    RequestProfiler.reset()
    t_start = time.time()

    logger.info("Executing Remediation Agent...")
    
    regulation_id_str = state.get("regulation_id")
    organization_id_str = state.get("organization_id")
    # Retrieve matching document list from state (or default to empty if not relevant)
    matched_doc_ids = state.get("matched_document_ids", [])
    
    if not regulation_id_str or not organization_id_str:
        logger.error("Regulation ID or Organization ID missing in state.")
        return {"remediation_draft_ids": []}
        
    db = SessionLocal()
    draft_ids = []
    
    try:
        from app.core.exceptions import LLMConfigurationError
        
        regulation = db.query(RegulationUpdate).filter(
            RegulationUpdate.id == uuid.UUID(regulation_id_str)
        ).first()
        if not regulation:
            logger.error("Regulation not found in database.")
            return {"remediation_draft_ids": []}

        if not matched_doc_ids:
            logger.info("No matched document IDs found. Generating a Proposed New SOP from scratch...")
            draft_ids = _generate_proposed_new_sop(db, regulation, organization_id_str)
            return {"remediation_draft_ids": draft_ids}

        # Read system prompt
        current_dir = os.path.dirname(os.path.abspath(__file__))
        prompt_path = os.path.join(current_dir, "../prompts/remediation.md")
        try:
            with open(prompt_path, "r", encoding="utf-8") as f:
                system_prompt = f.read()
        except Exception:
            system_prompt = (
                "You are a GxP compliance officer. Review the SOP and propose a revised "
                "version to comply with the regulation. Cite the exact clause numbers."
            )

        # Process matched documents in parallel
        import concurrent.futures
        
        max_workers = min(len(matched_doc_ids), 4)
        logger.info(f"Processing {len(matched_doc_ids)} matched documents in parallel using {max_workers} threads...")
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {
                executor.submit(
                    _process_single_document,
                    doc_id_val,
                    regulation_id_str,
                    organization_id_str,
                    system_prompt
                ): doc_id_val
                for doc_id_val in matched_doc_ids
            }
            
            for future in concurrent.futures.as_completed(futures):
                doc_val = futures[future]
                try:
                    draft_id, thread_metrics = future.result()
                    if draft_id:
                        draft_ids.append(draft_id)
                    # Aggregate thread-local metrics into the main profiler
                    for metric, val in thread_metrics.items():
                        if metric != "total_time":
                            RequestProfiler.log_metric(metric, val)
                except Exception as exc:
                    logger.error(f"Document thread processing for {doc_val} raised exception: {exc}")

        if not draft_ids and matched_doc_ids:
            logger.warning(
                "All matched documents were skipped or failed; falling back to Proposed New SOP generation."
            )
            draft_ids = _generate_proposed_new_sop(db, regulation, organization_id_str)

        logger.info(f"Remediation Agent successfully generated {len(draft_ids)} drafts.")
        RequestProfiler.log_metric("total_time", time.time() - t_start)
        RequestProfiler.print_summary("Remediation Agent")
        return {"remediation_draft_ids": draft_ids}
        
    except LLMConfigurationError:
        logger.error("Remediation Agent failed due to LLM configuration error.")
        db.rollback()
        raise
    except Exception as e:
        logger.error(f"Remediation Agent failed: {e}")
        db.rollback()
        if not llm_client.is_offline_mode():
            raise e
        return {"remediation_draft_ids": []}
    finally:
        db.close()
