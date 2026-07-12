import os
import shutil
import difflib
import re
import fitz  # PyMuPDF
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def find_section_number(lines, index):
    # Look back up to 30 lines to find nearest section header
    for idx in range(min(index, len(lines) - 1), max(-1, index - 30), -1):
        line = lines[idx].strip()
        # Match common section patterns like "1.2 ", "Section 2.1", "4.0.1"
        match = re.match(r'^((?:Section\s+)?\d+(?:\.\d+)*)', line, re.IGNORECASE)
        if match:
            return match.group(1)
    return None


def _rect_key(rect, precision=1):
    """Round rect coords for geometric deduplication."""
    return (
        round(rect.x0, precision),
        round(rect.y0, precision),
        round(rect.x1, precision),
        round(rect.y1, precision),
    )


def _dedupe_rects(rects):
    """Keep unique rects by rounded geometry (first occurrence wins)."""
    seen = set()
    unique = []
    for rect in rects:
        key = _rect_key(rect)
        if key in seen:
            continue
        seen.add(key)
        unique.append(rect)
    return unique


def _search_line_rects(page, line):
    """Exact text search only — avoid short-substring false positives."""
    if not line:
        return []
    return list(page.search_for(line)) or []


def _action_fingerprint(title, section, orig_text, proposed_text, justification, regulation_reference):
    return (
        title,
        section or "N/A",
        orig_text or "",
        proposed_text or "",
        justification or "",
        regulation_reference or "",
    )


def _add_highlight(page, rects, title, annot_content, stroke):
    """Write one content-bearing highlight covering all rects on this page."""
    unique_rects = _dedupe_rects(rects)
    if not unique_rects:
        return
    # PyMuPDF accepts a single Rect or a list of quads/rects for one annot
    annot = page.add_highlight_annot(unique_rects if len(unique_rects) > 1 else unique_rects[0])
    annot.set_colors(stroke=stroke)
    annot.set_info(title=title, content=annot_content)
    annot.update()


def apply_remediation(
    original_file_path: str,
    new_file_path: str,
    proposed_text: str,
    diff_content: dict,
    justification: str = None,
    regulation_reference: str = None
) -> None:
    _, ext = os.path.splitext(original_file_path.lower())
    
    # Read original text for diff alignment
    original_text = ""
    if ext == ".txt":
        try:
            with open(original_file_path, "r", encoding="utf-8") as f:
                original_text = f.read()
        except Exception:
            pass
    elif ext == ".docx":
        try:
            doc = docx.Document(original_file_path)
            original_text = "\n".join([p.text for p in doc.paragraphs])
        except Exception:
            pass
    elif ext == ".pdf":
        try:
            pdf_doc = fitz.open(original_file_path)
            original_text = ""
            for page in pdf_doc:
                original_text += page.get_text()
            pdf_doc.close()
        except Exception:
            pass

    if not original_text:
        # Fallback to diff_content's current_content if original text read failed
        original_text = diff_content.get("current_content") or ""

    if ext == ".txt":
        with open(new_file_path, "w", encoding="utf-8") as f:
            f.write(proposed_text)
            
    elif ext == ".docx":
        try:
            # 1. Copy the original document to the destination first.
            # This guarantees we NEVER start with an empty document and preserve all headers, footers, logos, tables, layout.
            shutil.copy(original_file_path, new_file_path)
            
            # 2. Open the copied file to perform in-place replacement
            doc = docx.Document(new_file_path)
            
            orig_paras = [p.text for p in doc.paragraphs]
            prop_paras = proposed_text.splitlines()
            
            # Align original paragraphs with proposed paragraphs using SequenceMatcher
            sm = difflib.SequenceMatcher(None, orig_paras, prop_paras)
            opcodes = sm.get_opcodes()
            
            # Apply edits in reverse order to prevent index shifting issues
            for tag, i1, i2, j1, j2 in reversed(opcodes):
                if tag == 'equal':
                    continue
                elif tag == 'replace':
                    num_orig = i2 - i1
                    num_prop = j2 - j1
                    for k in range(min(num_orig, num_prop)):
                        doc.paragraphs[i1 + k].text = prop_paras[j1 + k]
                    
                    if num_orig < num_prop:
                        # Insert remaining proposed paragraphs in-place
                        ref_para = doc.paragraphs[i2] if i2 < len(doc.paragraphs) else None
                        for k in range(num_orig, num_prop):
                            val = prop_paras[j1 + k]
                            if ref_para:
                                ref_para.insert_paragraph_before(val)
                            else:
                                doc.add_paragraph(val)
                    elif num_orig > num_prop:
                        # Remove extra original paragraphs in reverse to prevent shifting
                        for k in reversed(range(num_prop, num_orig)):
                            p = doc.paragraphs[i1 + k]
                            p_element = p._element
                            p_element.getparent().remove(p_element)
                            
                elif tag == 'delete':
                    # Delete extra paragraphs in reverse to prevent index shifting
                    for idx in reversed(range(i1, i2)):
                        p = doc.paragraphs[idx]
                        p_element = p._element
                        p_element.getparent().remove(p_element)
                        
                elif tag == 'insert':
                    ref_para = doc.paragraphs[i1] if i1 < len(doc.paragraphs) else None
                    for idx in range(j1, j2):
                        val = prop_paras[idx]
                        if ref_para:
                            ref_para.insert_paragraph_before(val)
                        else:
                            doc.add_paragraph(val)
                            
            doc.save(new_file_path)
        except Exception as e:
            # Fallback: if in-place editing crashes for any reason, we already copied the original file
            # to new_file_path, so formatting is preserved. We do NOT overwrite it with a blank document.
            print(f"Error during DOCX template editing: {e}")
            
    elif ext == ".pdf":
        try:
            # For PDF, do NOT modify the document text. Instead, generate an annotated review PDF.
            # Build canonical unique remediation actions first, then write one highlight
            # annotation per action per page (multi-quad), never one card per rect.
            doc = fitz.open(original_file_path)

            original_lines = original_text.splitlines()
            proposed_lines = proposed_text.splitlines()

            just = justification if justification else "AI Compliance Recommendation"
            reg_ref = regulation_reference if regulation_reference else "FDA Regulation"

            # fingerprint -> {title, section, orig, proposed, justification, regulation, pages: {page_idx: [rects]}}
            actions = {}

            sm = difflib.SequenceMatcher(None, original_lines, proposed_lines)
            for tag, i1, i2, j1, j2 in sm.get_opcodes():
                if tag == "equal":
                    continue

                rem_lines = [line.strip() for line in original_lines[i1:i2] if line.strip()]
                add_text = "\n".join(proposed_lines[j1:j2]).strip()
                section_num = find_section_number(original_lines, i1) or "N/A"
                orig_text_block = "\n".join(rem_lines) if rem_lines else "N/A (Insertion)"
                proposed_block = add_text if add_text else "N/A (Deletion)"

                if tag == "delete":
                    title = "AI Remediation - Deletion Suggestion"
                    stroke = (0.9, 0.2, 0.2)
                elif tag == "insert":
                    title = "AI Remediation - Insertion Suggestion"
                    stroke = (0.2, 0.8, 0.2)
                else:
                    title = "AI Remediation - Replacement Suggestion"
                    stroke = (1.0, 0.9, 0.0)

                fp = _action_fingerprint(
                    title, section_num, orig_text_block, proposed_block, just, reg_ref
                )
                if fp not in actions:
                    actions[fp] = {
                        "title": title,
                        "section": section_num,
                        "orig_text": orig_text_block,
                        "proposed_text": proposed_block,
                        "justification": just,
                        "regulation_reference": reg_ref,
                        "stroke": stroke,
                        "pages": {},
                    }

                page_rects = actions[fp]["pages"]

                if tag == "insert":
                    anchor_idx = i1 - 1
                    anchor_line = None
                    while anchor_idx >= 0:
                        if original_lines[anchor_idx].strip():
                            anchor_line = original_lines[anchor_idx].strip()
                            break
                        anchor_idx -= 1

                    if anchor_line:
                        for page_idx, page in enumerate(doc):
                            rects = _search_line_rects(page, anchor_line)
                            if rects:
                                # One anchor rect is enough for insertion markers
                                page_rects.setdefault(page_idx, []).append(rects[0])
                                break
                    continue

                for page_idx, page in enumerate(doc):
                    rects = []
                    for line in rem_lines:
                        # Exact match only; take first hit per line to avoid boilerplate fan-out
                        line_rects = _search_line_rects(page, line)
                        if line_rects:
                            rects.append(line_rects[0])
                    if rects:
                        page_rects.setdefault(page_idx, []).extend(rects)

            for action in actions.values():
                annot_content = (
                    f"Section: {action['section']}\n"
                    f"Original Text: {action['orig_text']}\n"
                    f"Proposed Text: {action['proposed_text']}\n"
                    f"Justification: {action['justification']}\n"
                    f"Regulation Reference: {action['regulation_reference']}"
                )
                for page_idx, rects in action["pages"].items():
                    _add_highlight(
                        doc[page_idx],
                        rects,
                        action["title"],
                        annot_content,
                        action["stroke"],
                    )

            doc.save(new_file_path, incremental=False, encryption=fitz.PDF_ENCRYPT_KEEP)
            doc.close()
        except Exception as e:
            # Fallback: copy original if anything fails
            print(f"Error during PDF annotation: {e}")
            try:
                shutil.copy(original_file_path, new_file_path)
            except Exception:
                pass
    else:
        # Save other files directly
        with open(new_file_path, "w", encoding="utf-8") as f:
            f.write(proposed_text)
