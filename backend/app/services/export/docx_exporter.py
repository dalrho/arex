import io
from datetime import datetime
import docx
from docx.shared import Pt, RGBColor
from app.models.remediation_draft import RemediationDraft
from app.models.document import Document
from app.models.regulation_update import RegulationUpdate

def generate_docx_report(draft: RemediationDraft, doc: Document, reg: RegulationUpdate) -> bytes:
    """
    Generates a DOCX document with additions/deletions highlighting using python-docx.
    """
    document = docx.Document()
    
    # Title Section
    title = document.add_heading("SENTINEL OS - REGULATORY COMPLIANCE EXPORT", level=0)
    
    # Metadata Block
    p = document.add_paragraph()
    p.add_run("Source SOP Document: ").bold = True
    p.add_run(f"{doc.filename}\n")
    p.add_run("Active SOP Version: ").bold = True
    p.add_run(f"v{doc.version}\n")
    p.add_run("FDA Regulation Title: ").bold = True
    p.add_run(f"{reg.title}\n")
    
    p.add_run("Remediation Status: ").bold = True
    status_run = p.add_run(draft.status)
    status_run.bold = True
    if draft.status == "APPROVED":
        status_run.font.color.rgb = RGBColor(47, 133, 90) # Green
    else:
        status_run.font.color.rgb = RGBColor(197, 48, 48) # Red
        
    p.add_run(f"\nExported At: ").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    
    # Section 1: Additions
    document.add_heading("PROPOSED COMPLIANCE ADDITIONS (+)", level=1)
    added_lines = draft.diff_content.get("added", []) if draft.diff_content else []
    if not added_lines:
        p_empty = document.add_paragraph("No direct additions proposed.")
        p_empty.runs[0].font.color.rgb = RGBColor(113, 128, 150) # Gray
    else:
        for line in added_lines:
            p_add = document.add_paragraph()
            run = p_add.add_run(f"+ {line}")
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(47, 133, 90) # Green
            
    # Section 2: Deletions
    document.add_heading("PROPOSED COMPLIANCE DELETIONS (-)", level=1)
    removed_lines = draft.diff_content.get("removed", []) if draft.diff_content else []
    if not removed_lines:
        p_empty = document.add_paragraph("No direct deletions proposed.")
        p_empty.runs[0].font.color.rgb = RGBColor(113, 128, 150) # Gray
    else:
        for line in removed_lines:
            p_rem = document.add_paragraph()
            run = p_rem.add_run(f"- {line}")
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(197, 48, 48) # Red

    file_stream = io.BytesIO()
    document.save(file_stream)
    docx_bytes = file_stream.getvalue()
    file_stream.close()
    return docx_bytes
