import os
import tempfile
import docx
import fitz
import pytest
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from app.services.document_modifier import apply_remediation

def test_docx_remediation():
    # 1. Create a sample DOCX with some structure
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "test.docx")
        new_doc_path = os.path.join(tmpdir, "test_v2.docx")
        
        doc = docx.Document()
        doc.add_heading("Standard Operating Procedure", 0)
        p1 = doc.add_paragraph("This is the original paragraph one. It is very standard.")
        p2 = doc.add_paragraph("This is paragraph two. It should not be modified.")
        doc.save(doc_path)
        
        original_text = "Standard Operating Procedure\nThis is the original paragraph one. It is very standard.\nThis is paragraph two. It should not be modified."
        proposed_text = "Standard Operating Procedure\nThis is the revised paragraph one. It is updated.\nThis is paragraph two. It should not be modified."
        
        diff_content = {
            "current_content": original_text,
            "removed": ["This is the original paragraph one. It is very standard."],
            "added": ["This is the revised paragraph one. It is updated."]
        }
        
        # 2. Apply remediation
        apply_remediation(doc_path, new_doc_path, proposed_text, diff_content)
        
        # 3. Verify
        assert os.path.exists(new_doc_path)
        new_doc = docx.Document(new_doc_path)
        
        # Check headings and paragraphs count
        assert len(new_doc.paragraphs) == len(doc.paragraphs)
        assert new_doc.paragraphs[0].text == "Standard Operating Procedure"
        assert new_doc.paragraphs[1].text == "This is the revised paragraph one. It is updated."
        assert new_doc.paragraphs[2].text == "This is paragraph two. It should not be modified."

def test_docx_formatting_preservation():
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "formatting_test.docx")
        new_doc_path = os.path.join(tmpdir, "formatting_test_v2.docx")
        
        doc = docx.Document()
        
        # Add header and footer
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = "CONFIDENTIAL - SOP HEADER"
        footer = section.footer
        footer.paragraphs[0].text = "Page 1 of 1 - SOP FOOTER"
        
        # Add heading
        doc.add_heading("SOP Document Title", 0)
        
        # Add normal paragraph
        doc.add_paragraph("Original paragraph body.")
        
        # Add numbered list item
        doc.add_paragraph("First numbered item", style='List Number')
        
        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "Value A"
        table.cell(1, 1).text = "Value B"
        
        doc.save(doc_path)
        
        # original_text and proposed_text only contain the body paragraphs, matching documents.py parsed_text extraction.
        original_text = "SOP Document Title\nOriginal paragraph body.\nFirst numbered item"
        proposed_text = "SOP Document Title\nRevised paragraph body text here.\nFirst numbered item"
        
        diff_content = {
            "current_content": original_text,
            "removed": ["Original paragraph body."],
            "added": ["Revised paragraph body text here."]
        }
        
        apply_remediation(doc_path, new_doc_path, proposed_text, diff_content)
        
        # Verify formatting remains intact
        assert os.path.exists(new_doc_path)
        new_doc = docx.Document(new_doc_path)
        
        # Verify header and footer are preserved
        new_section = new_doc.sections[0]
        assert new_section.header.paragraphs[0].text == "CONFIDENTIAL - SOP HEADER"
        assert new_section.footer.paragraphs[0].text == "Page 1 of 1 - SOP FOOTER"
        
        # Verify table elements and structure are preserved exactly
        assert len(new_doc.tables) == 1
        new_table = new_doc.tables[0]
        assert len(new_table.rows) == 2
        assert len(new_table.columns) == 2
        assert new_table.cell(0, 0).text == "Header A"
        assert new_table.cell(0, 1).text == "Header B"
        assert new_table.cell(1, 0).text == "Value A"
        assert new_table.cell(1, 1).text == "Value B"
        
        # Verify body paragraphs are correctly edited in-place
        assert len(new_doc.paragraphs) == 3
        assert new_doc.paragraphs[0].text == "SOP Document Title"
        assert new_doc.paragraphs[1].text == "Revised paragraph body text here."
        assert new_doc.paragraphs[2].text == "First numbered item"
        assert new_doc.paragraphs[2].style.name == 'List Number'

def test_pdf_remediation():
    # 1. Create a sample PDF
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "test.pdf")
        new_pdf_path = os.path.join(tmpdir, "test_v2.pdf")
        
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = [
            Paragraph("Standard Operating Procedure", styles['Heading1']),
            Paragraph("This is the original text segment.", styles['Normal']),
            Paragraph("This is other text that remains the same.", styles['Normal'])
        ]
        doc.build(story)
        
        original_text = "Standard Operating Procedure\nThis is the original text segment.\nThis is other text that remains the same."
        proposed_text = "Standard Operating Procedure\nThis is the new revised text segment.\nThis is other text that remains the same."
        
        diff_content = {
            "current_content": original_text,
            "removed": ["This is the original text segment."],
            "added": ["This is the new revised text segment."]
        }
        
        # 2. Apply remediation
        apply_remediation(
            original_file_path=pdf_path,
            new_file_path=new_pdf_path,
            proposed_text=proposed_text,
            diff_content=diff_content,
            justification="Test Justification Detail",
            regulation_reference="21 CFR Part 11"
        )
        
        # 3. Verify PDF exists and has the same page count (1 page)
        assert os.path.exists(new_pdf_path)
        pdf_doc = fitz.open(new_pdf_path)
        assert pdf_doc.page_count == 1
        
        # Verify the original text is NOT removed or modified in the PDF layout
        page = pdf_doc[0]  # Store in local variable to prevent GC of the page object
        page_text = page.get_text()
        assert "This is the original text segment." in page_text
        assert "This is the new revised text segment." not in page_text
        assert "This is other text that remains the same." in page_text
        
        # Verify PDF contains the review highlight annotation with the suggestions
        annots = list(page.annots())
        assert len(annots) >= 1
        annot = annots[0]
        assert annot.type[1] == "Highlight"
        
        # Colors match (yellow highlight for replacement suggestion)
        stroke = annot.colors["stroke"]
        assert stroke[0] == pytest.approx(1.0)
        assert stroke[1] == pytest.approx(0.9)
        assert stroke[2] == pytest.approx(0.0)
        
        # Info dictionary contains remediation metadata
        info = annot.info
        assert info["title"] == "AI Remediation - Replacement Suggestion"
        assert "Original Text: This is the original text segment." in info["content"]
        assert "Proposed Text: This is the new revised text segment." in info["content"]
        assert "Justification: Test Justification Detail" in info["content"]
        assert "Regulation Reference: 21 CFR Part 11" in info["content"]
        
        pdf_doc.close()


def test_pdf_multiline_replace_creates_one_annotation():
    """Multi-line replace must produce one content-bearing annot, not one per line/rect."""
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "multi.pdf")
        new_pdf_path = os.path.join(tmpdir, "multi_v2.pdf")

        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = [
            Paragraph("1.0 CAPA Procedure", styles["Heading1"]),
            Paragraph("Line alpha requires update for compliance.", styles["Normal"]),
            Paragraph("Line beta requires update for compliance.", styles["Normal"]),
            Paragraph("Line gamma requires update for compliance.", styles["Normal"]),
            Paragraph("Unchanged trailing paragraph stays put.", styles["Normal"]),
        ]
        doc.build(story)

        original_text = (
            "1.0 CAPA Procedure\n"
            "Line alpha requires update for compliance.\n"
            "Line beta requires update for compliance.\n"
            "Line gamma requires update for compliance.\n"
            "Unchanged trailing paragraph stays put."
        )
        proposed_text = (
            "1.0 CAPA Procedure\n"
            "Line alpha has been revised for compliance.\n"
            "Line beta has been revised for compliance.\n"
            "Line gamma has been revised for compliance.\n"
            "Unchanged trailing paragraph stays put."
        )

        apply_remediation(
            original_file_path=pdf_path,
            new_file_path=new_pdf_path,
            proposed_text=proposed_text,
            diff_content={"current_content": original_text},
            justification="CAPA alignment",
            regulation_reference="21 CFR 820",
        )

        pdf_doc = fitz.open(new_pdf_path)
        page = pdf_doc[0]
        annots = list(page.annots() or [])
        # One logical remediation action → one annotation (multi-quad OK)
        assert len(annots) == 1
        content = annots[0].info["content"]
        assert "Line alpha requires update for compliance." in content
        assert "Line beta requires update for compliance." in content
        assert "Line gamma requires update for compliance." in content
        assert "Line alpha has been revised for compliance." in content
        pdf_doc.close()


def test_pdf_repeated_boilerplate_does_not_fan_out():
    """Repeated identical lines must not create one annotation per occurrence."""
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "boilerplate.pdf")
        new_pdf_path = os.path.join(tmpdir, "boilerplate_v2.pdf")

        repeated = "CONFIDENTIAL - Do not distribute without approval."
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = [
            Paragraph("2.0 Document Control", styles["Heading1"]),
            Paragraph(repeated, styles["Normal"]),
            Paragraph("Unique body paragraph that will change.", styles["Normal"]),
            Paragraph(repeated, styles["Normal"]),
            Paragraph(repeated, styles["Normal"]),
        ]
        doc.build(story)

        original_text = (
            f"2.0 Document Control\n"
            f"{repeated}\n"
            f"Unique body paragraph that will change.\n"
            f"{repeated}\n"
            f"{repeated}"
        )
        proposed_text = (
            f"2.0 Document Control\n"
            f"{repeated}\n"
            f"Unique body paragraph that was revised.\n"
            f"{repeated}\n"
            f"{repeated}"
        )

        apply_remediation(
            original_file_path=pdf_path,
            new_file_path=new_pdf_path,
            proposed_text=proposed_text,
            diff_content={"current_content": original_text},
            justification="Body update only",
            regulation_reference="ISO 13485",
        )

        pdf_doc = fitz.open(new_pdf_path)
        all_annots = []
        for page in pdf_doc:
            page_annots = list(page.annots() or [])
            all_annots.extend(page_annots)
        # Only the unique body change should produce a recommendation annot
        assert len(all_annots) == 1
        assert "Unique body paragraph that will change." in all_annots[0].info["content"]
        assert "Unique body paragraph that was revised." in all_annots[0].info["content"]
        pdf_doc.close()
